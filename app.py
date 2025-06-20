# =============================================================================
# PDF to DOCX/TXT Converter with OCR v4.2.0 Enterprise Edition
# Author: Alan Steinbarth
# Date: 20 czerwca 2025
# =============================================================================

"""
# 📋 PDF to DOCX/TXT Converter - Dokumentacja Techniczna

## 🎯 Opis Aplikacji
Zaawansowany, wieloplatformowy konwerter PDF na DOCX/TXT z automatycznym 
rozpoznawaniem tekstu (OCR). Aplikacja oferuje nowoczesny interfejs GUI,
wsparcie dla motywów jasny/ciemny oraz enterprise-grade funkcjonalności.

## 📑 Spis Treści

### 1. 📦 Importy i Konfiguracja (linie 50-80)
- Importy standardowych bibliotek Python
- Importy bibliotek zewnętrznych (PyMuPDF, PIL, pdfminer)
- Konfiguracja logowania aplikacji

### 2. 🖥️ Obsługa Logów GUI (linie 85-100)
- `GuiLogHandler`: Custom handler dla wyświetlania logów w GUI

### 3. ⚙️ Funkcje Pomocnicze (linie 105-200)
- `load_config()`: Ładowanie konfiguracji z YAML
- `preprocess_image()`: Preprocessing obrazów dla OCR

### 4. 🏗️ Klasa Główna PDFConverterApp (linie 205-1200)

#### 4.1 Inicjalizacja i Setup
- `__init__()`: Inicjalizacja aplikacji
- `get_default_output_dir()`: Określanie domyślnego folderu wyjściowego

#### 4.2 System Motywów
- `detect_system_theme()`: Wykrywanie motywu systemowego
- `setup_theme_colors()`: Konfiguracja kolorów motywów
- `toggle_theme()` / `update_theme()`: Przełączanie motywów

#### 4.3 Interface Użytkownika
- `create_menu()`: Tworzenie paska menu
- `show_about_dialog()`: Dialog "O programie"
- `create_widgets()`: Budowanie głównego UI

#### 4.4 Obsługa Plików
- `browse_files()` / `browse_folder()`: Wybór plików/folderów
- `_process_single_file()`: Konwersja pojedynczego pliku

#### 4.5 Konwersja i OCR
- `_perform_conversion_threaded()`: Główna logika konwersji
- `_extract_text_with_ocr()`: Rozpoznawanie tekstu z OCR

#### 4.6 Podgląd PDF
- `_show_pdf_preview()`: Wyświetlanie podglądu PDF
- `_get_pixmap_compat()`: Kompatybilność z różnymi wersjami PyMuPDF

### 5. 🚀 Punkt Wejścia (linia 1200+)
- Uruchomienie aplikacji

## 🔧 Kluczowe Technologie
- **GUI**: Tkinter z custom styling
- **PDF Processing**: PyMuPDF, pdfminer
- **OCR**: Tesseract (opcjonalne)
- **Documents**: python-docx
- **Configuration**: PyYAML
- **Testing**: pytest

## 🌟 Enterprise Features
- ✅ Logowanie do pliku z rotacją
- ✅ Konfiguracja przez YAML
- ✅ Testy jednostkowe
- ✅ CI/CD workflow
- ✅ Cross-platform support
- ✅ Error handling
- ✅ Progress tracking
"""

# =============================================================================
# 📦 IMPORTY I KONFIGURACJA
# =============================================================================

# Biblioteki standardowe Python
import logging
import logging.handlers
import os
import platform
import subprocess
import sys
import threading
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, Toplevel, messagebox
import traceback

# Biblioteki zewnętrzne
import fitz  # PyMuPDF - przetwarzanie PDF
from PIL import Image, ImageTk  # Pillow - przetwarzanie obrazów
from pdfminer.high_level import extract_text, extract_pages  # pdfminer - ekstrakcja tekstu
import yaml  # PyYAML - konfiguracja

# =============================================================================
# 📝 KONFIGURACJA LOGOWANIA I TESSERACT
# =============================================================================

# Konfiguracja Tesseract OCR
try:
    import pytesseract as pt_module
    if platform.system() == "Darwin":  # macOS
        tesseract_paths = [
            "/opt/homebrew/bin/tesseract",
            "/usr/local/bin/tesseract"
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pt_module.pytesseract.tesseract_cmd = path
                break
    pytesseract = pt_module
    TESSERACT_AVAILABLE = True
    logging.info("Tesseract OCR został pomyślnie zainicjowany.")
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("Pytesseract nie jest zainstalowany. Funkcje OCR będą niedostępne.")
    # raise  # USUNIĘTO raise, aplikacja działa bez OCR
except (OSError, AttributeError, RuntimeError) as e:
    TESSERACT_AVAILABLE = False
    logging.error("Błąd podczas inicjalizacji Tesseract: %s", e)
    # raise  # USUNIĘTO raise, aplikacja działa bez OCR

# Ustawienie zmiennej środowiskowej dla macOS - ZAWSZE na początku
if platform.system() == "Darwin":
    os.environ['TK_SILENCE_DEPRECATION'] = '1'

# =============================================================================
# 🖥️ OBSŁUGA LOGÓW W GUI
# =============================================================================

class GuiLogHandler(logging.Handler):
    """
    ## 📋 Custom Logging Handler for GUI
    
    **Opis**: Niestandardowy handler logowania, który wyświetla komunikaty 
    logów bezpośrednio w widget tekstowym GUI aplikacji.
    
    **Funkcjonalność**:
    - Przekierowuje logi do Text widget w interfejsie
    - Formatuje komunikaty z timestamp
    - Automatyczne przewijanie do najnowszych wpisów
    
    **Parametry**:
    - `text_widget`: tkinter.Text widget do wyświetlania logów
    """
    
    def __init__(self, text_widget):
        super().__init__()
        self.text_widget = text_widget
        self.formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s', datefmt='%H:%M:%S')
        self.setFormatter(self.formatter)

    def emit(self, record):
        msg = self.format(record)
        def append():
            self.text_widget.config(state=tk.NORMAL)
            self.text_widget.insert(tk.END, msg + "\n")  # Każdy log w osobnej linii
            self.text_widget.config(state=tk.DISABLED)
            self.text_widget.see(tk.END)
        if hasattr(self.text_widget, 'after'):
            self.text_widget.after(0, append)

# Inicjalizacja głównego loggera
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)

# --- Enterprise: logowanie do pliku z rotacją ---
LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
if not os.path.exists(LOG_DIR):
    os.makedirs(LOG_DIR)
LOG_FILE = os.path.join(LOG_DIR, 'app.log')
file_handler = logging.handlers.RotatingFileHandler(
    LOG_FILE, maxBytes=2*1024*1024, backupCount=5, encoding='utf-8'
)
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s', datefmt='%Y-%m-%d %H:%M:%S'))
logging.getLogger().addHandler(file_handler)

# =============================================================================
# ⚙️ FUNKCJE POMOCNICZE I KONFIGURACJA
# =============================================================================

def load_config():
    """
    ## 📁 Ładowanie Konfiguracji z YAML
    
    **Opis**: Wczytuje konfigurację aplikacji z pliku `config.yaml`.
    W przypadku braku pliku lub błędu, używa wartości domyślnych.
    
    **Lokalizacja pliku**: `./config.yaml` (katalog roboczy)
    
    **Dostępne opcje**:
    - `output_dir`: Domyślny folder wyjściowy
    - `log_level`: Poziom logowania (DEBUG, INFO, WARNING, ERROR)
    - `ocr_language`: Język OCR (domyślnie 'pol')
    - `ocr_dpi`: DPI dla OCR (domyślnie 300)
    - `ocr_psm`: Page Segmentation Mode (domyślnie 6)
    
    **Returns**:
        dict: Słownik z konfiguracją aplikacji
        
    **Example**:
        ```yaml
        output_dir: "~/Desktop"
        log_level: "INFO"
        ocr_language: "pol"
        ```
    """
    # Szukaj config.yaml w bieżącym katalogu roboczym (dla testów i uruchomień enterprise)
    config_path = os.path.join(os.getcwd(), 'config.yaml')
    default = {
        'output_dir': os.path.expanduser('~/Desktop'),
        'log_level': 'INFO',
        'ocr_lang': 'pol',
    }
    if not os.path.exists(config_path):
        return default
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            cfg = yaml.safe_load(f)
        if not isinstance(cfg, dict):
            return default
        merged = default.copy()
        merged.update({k: v for k, v in cfg.items() if v is not None})
        merged['output_dir'] = os.path.expanduser(merged['output_dir'])
        return merged
    except (OSError, IOError, yaml.YAMLError, KeyError, TypeError) as e:
        logging.warning("Nie udało się wczytać config.yaml: %s", e)
        return default

# =============================================================================
# 🏗️ KLASA GŁÓWNA APLIKACJI
# =============================================================================

class PDFtoDocxConverterApp(tk.Tk):
    """
    ## 🎯 Główna Klasa Aplikacji PDF Converter
    
    **Opis**: Główna klasa aplikacji konwertera PDF do DOCX/TXT z obsługą OCR.
    Dziedziczy po `tkinter.Tk` i implementuje pełny interfejs użytkownika.
    
    **Główne funkcjonalności**:
    - 📄 Konwersja PDF → DOCX/TXT
    - 🔍 OCR dla zeskanowanych dokumentów
    - 🎨 System motywów jasny/ciemny
    - 📊 Progress tracking i logowanie
    - 🖼️ Podgląd PDF
    - ⚙️ Konfiguracja przez YAML
    - 🔄 Batch processing
    
    **Obsługiwane platformy**: macOS, Windows, Linux
    
    **Architektura**:
    - Model-View z separacją logiki biznesowej
    - Threading dla operacji długotrwałych
    - Event-driven GUI updates
    - Error handling z graceful degradation
    """
    
    def get_default_output_dir(self):
        """
        ## 📁 Określanie Domyślnego Folderu Wyjściowego
        
        **Opis**: Automatycznie wykrywa i ustawia najbardziej odpowiedni 
        folder wyjściowy na podstawie systemu operacyjnego.
        
        **Logika wyboru**:
        
        ### Windows:
        1. `~/OneDrive/Pulpit` (jeśli istnieje)
        2. `~/Desktop` (fallback)
        3. `~` (ostateczny fallback)
        
        ### macOS:
        1. `~/Desktop` (standardowy pulpit)
        2. `~` (fallback)
        
        ### Linux:
        1. `~/Desktop` (standardowy pulpit)
        2. `~` (fallback)
        
        **Returns**:
            str: Ścieżka do domyślnego folderu wyjściowego
            
        **Obsługa błędów**: Graceful fallback do katalogu domowego użytkownika
        """
        system = platform.system()
        try:
            if system == "Windows":
                onedrive_desktop = os.path.expanduser("~/OneDrive/Pulpit")
                local_desktop = os.path.expanduser("~/Desktop")
                if os.path.exists(onedrive_desktop):
                    return onedrive_desktop
                if os.path.exists(local_desktop):
                    return local_desktop
                return os.path.expanduser("~")
            if system == "Darwin":  # macOS
                desktop = os.path.expanduser("~/Desktop")
                return desktop if os.path.exists(desktop) else os.path.expanduser("~")
            # Linux i inne systemy Unix
            desktop = os.path.expanduser("~/Desktop")
            return desktop if os.path.exists(desktop) else os.path.expanduser("~")
        except (OSError, IOError):
            return os.path.expanduser("~")

    def _bring_to_front_macos(self):
        """Wymusza pokazanie okna na wierzchu na macOS po pełnej inicjalizacji (silniejsze)."""
        self.deiconify()
        self.lift()
        self.focus_force()
        self.attributes('-topmost', True)
        self.after(500, lambda: self.attributes('-topmost', False))
        # Dodatkowo wymuś aktywację przez AppleScript (jeśli dostępne)
        try:
            subprocess.Popen([
                'osascript', '-e',
                'tell application "System Events" to set frontmost of the first process whose unix id is (do shell script "echo $PPID") to true'
            ])
        except (OSError, subprocess.SubprocessError):
            pass

    def __init__(self):
        self.theme_switch_btn = None  # Inicjalizacja atrybutu przed create_widgets
        self.theme_mode = self.detect_system_theme()  # light lub dark
        self.theme_colors = {
            "light": {
                "bg": "#f5f5f7",
                "panel": "#ffffff",
                "listbox": "#ffffff",
                "label": "#1d1d1f",
                "button": "#007aff",
                "button_text": "#ffffff",
                "button_active": "#0056cc",
                "button_secondary": "#f2f2f7",
                "button_secondary_text": "#007aff",
                "button_secondary_active": "#e5e5ea",
                "entry": "#ffffff",
                "entry_border": "#d2d2d7",
                "footer": "#f5f5f7",
                "footer_fg": "#86868b",
                "frame": "#f5f5f7",
                "progress": "#007aff",
                "progress_trough": "#e5e5ea",
                "border": "#d2d2d7",
                "accent": "#007aff"
            },
            "dark": {
                "bg": "#1c1c1e",
                "panel": "#2c2c2e",
                "listbox": "#2c2c2e",
                "label": "#ffffff",
                "button": "#0a84ff",
                "button_text": "#ffffff",
                "button_active": "#0066cc",
                "button_secondary": "#48484a",
                "button_secondary_text": "#0a84ff",
                "button_secondary_active": "#636366",
                "entry": "#2c2c2e",
                "entry_border": "#48484a",
                "footer": "#1c1c1e",
                "footer_fg": "#8e8e93",
                "frame": "#1c1c1e",
                "progress": "#0a84ff",
                "progress_trough": "#48484a",
                "border": "#48484a",
                "accent": "#0a84ff"
            }
        }
        self.app_config = load_config()
        logging.getLogger().setLevel(self.app_config.get('log_level', 'INFO'))
        super().__init__()
        # --- Diagnostyka widoczności widgetów na macOS ---
        if platform.system() == "Darwin":
            try:
                style = ttk.Style(self)
                style.theme_use('clam')  # 'clam' jest jasny i czytelny na macOS
                logging.info("Ustawiono styl ttk: clam (diagnostyka)")
            except (tk.TclError, AttributeError) as e:
                logging.warning("Nie udało się wymusić stylu ttk na macOS: %s", e)
        self.configure(bg="#f8f8f8")  # Jasne tło dla głównego okna
        # --- Poprawka: zawsze ustaw output_dir ---
        try:
            self.selected_files = []
            self.output_dir = self.app_config.get('output_dir') or self.get_default_output_dir()
        except (KeyError, TypeError, OSError) as e:
            logging.error("Błąd przy ustawianiu output_dir: %s", e)
            self.output_dir = self.get_default_output_dir()
        self.output_format = tk.StringVar(value="docx")
        self.conversion_thread = None
        self.cancel_event = threading.Event()
        self.current_preview_image = None
        self.ocr_lang = self.app_config.get('ocr_lang', 'pol')
        self.title("Zaawansowany Konwerter PDF v4.2.0 Enterprise Edition")
        self.geometry("1024x768")
        self.minsize(900, 600)
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        logging.info("Wywołanie create_widgets...")
        self.create_widgets()
        logging.info("create_widgets zakończone.")
        self._setup_logging()
        self.create_menu()
        self.center_window()
        self.update_idletasks()
        self.update()
        logging.info("Aplikacja uruchomiona.")
        self.on_file_select()
        # Tekst pomocy do okna 'O programie'
        self.help_text = (
            "Instrukcja obsługi:\n"
            "1. Dodaj pliki PDF do listy przyciskiem 'Wybierz pliki' lub 'Dodaj folder'.\n"
            "2. (Opcjonalnie) Zmień folder docelowy.\n"
            "3. Wybierz format wyjściowy (DOCX lub TXT).\n"
            "4. Kliknij 'Konwertuj', aby rozpocząć.\n"
            "5. Postęp i logi pojawią się na dole okna.\n"
            "\n"
            "Wymagane narzędzia zewnętrzne:\n"
            "- Poppler (do konwersji PDF na obrazy)\n"
            "- Tesseract OCR (do rozpoznawania tekstu ze skanów)\n"
            "\n"
            "Pobierz Poppler:\n"
            "  Windows: https://github.com/oschwartz10612/poppler-windows/releases\n"
            "  macOS: brew install poppler\n"
            "  Linux: sudo apt install poppler-utils\n"
            "\n"
            "Pobierz Tesseract OCR:\n"
            "  Windows: https://github.com/tesseract-ocr/tesseract\n"
            "  macOS: brew install tesseract tesseract-lang\n"
            "  Linux: sudo apt install tesseract-ocr tesseract-ocr-pol\n"
            "\n"
            "Więcej informacji i FAQ znajdziesz w pliku README.md."
        )
        # Wymuś pokazanie okna na wierzchu na macOS po pełnej inicjalizacji
        if platform.system() == "Darwin":
            self.after(200, self._bring_to_front_macos)

    # =========================================================================
    # 🎨 SYSTEM MOTYWÓW I KOLORÓW
    # =========================================================================

    def detect_system_theme(self):
        """
        ## 🎨 Wykrywanie Motywu Systemowego
        
        **Opis**: Automatycznie wykrywa motyw systemowy (jasny/ciemny) 
        na różnych platformach operacyjnych.
        
        **Obsługiwane platformy**:
        
        ### macOS:
        - Odczytuje `AppleInterfaceStyle` z defaults
        - "Dark" → motyw ciemny
        - Brak wartości → motyw jasny
        
        ### Windows:
        - Sprawdza registry `AppsUseLightTheme`
        - 0 → motyw ciemny
        - 1 → motyw jasny
        
        ### Linux:
        - Domyślnie motyw jasny (brak automatycznego wykrywania)
        
        **Returns**:
            str: "light" lub "dark"
            
        **Fallback**: W przypadku błędu zwraca "light"
        """
        try:
            if sys.platform == "darwin":
                result = subprocess.run([
                    'defaults', 'read', '-g', 'AppleInterfaceStyle'
                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                text=True, check=False)
                if 'Dark' in result.stdout:
                    return "dark"
            elif sys.platform.startswith("win"):
                try:
                    import winreg
                    reg_path = (r"Software\\Microsoft\\Windows\\CurrentVersion\\"
                               r"Themes\\Personalize")
                    with winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path) as key:
                        apps_use_light = winreg.QueryValueEx(key, "AppsUseLightTheme")[0]
                        return "light" if apps_use_light == 1 else "dark"
                except ImportError:
                    pass
        except (OSError, subprocess.SubprocessError):
            pass
        return "light"

    def _setup_logging(self):
        logger = logging.getLogger()
        logger.setLevel(logging.INFO)
        logging.getLogger('pdfminer').setLevel(logging.ERROR)
        logging.getLogger('PIL').setLevel(logging.WARNING)
        for handler in logger.handlers[:]:
            logger.removeHandler(handler)
        # Tworzenie widgetu logów jeśli nie istnieje
        if not hasattr(self, 'log_text_widget'):
            c = self.theme_colors[self.theme_mode]
            self.log_text_widget = scrolledtext.ScrolledText(
                self,
                height=8,
                state=tk.DISABLED,
                bg=c["listbox"],
                fg=c["label"],
                insertbackground=c["label"],
                selectbackground=c["accent"],
                selectforeground="white"
            )
            self.log_text_widget.pack(
                side=tk.BOTTOM, fill=tk.X, padx=10, pady=(0, 5)
            )
        gui_handler = GuiLogHandler(self.log_text_widget)
        logger.addHandler(gui_handler)

    def create_menu(self):
        """Tworzy pasek menu aplikacji."""
        self.menubar = tk.Menu(self)
        filemenu = tk.Menu(self.menubar, tearoff=0)
        filemenu.add_command(label="Wybierz pliki...", command=self.select_files, accelerator="Cmd+O" if platform.system() == "Darwin" else "Ctrl+O")
        filemenu.add_command(label="Dodaj folder...", command=self.add_folder, accelerator="Cmd+Shift+O" if platform.system() == "Darwin" else "Ctrl+Shift+O")
        filemenu.add_separator()
        filemenu.add_command(label="Wyczyść listę plików", command=self.clear_file_list)
        filemenu.add_separator()
        filemenu.add_command(label="Wyjdź", command=self.on_closing, accelerator="Cmd+Q" if platform.system() == "Darwin" else "Alt+F4")
        self.menubar.add_cascade(label="Plik", menu=filemenu)
        # Usunięto menu Opcje (OCR)
        helpmenu = tk.Menu(self.menubar, tearoff=0)
        helpmenu.add_command(label="O programie...", command=self.show_about_dialog)
        self.menubar.add_cascade(label="Pomoc", menu=helpmenu)
        self.config(menu=self.menubar)

    def show_about_dialog(self):
        """
        Display about dialog with application information.
        Wyświetla okno dialogowe 'O programie' oraz instrukcję i linki do Poppler/Tesseract.
        """
        about_win = Toplevel(self)
        about_win.title("O programie - PDF Converter v4.2.0 Enterprise Edition")
        dialog_width = 500
        dialog_height = 420  # zwiększona wysokość
        about_win.resizable(False, False)
        about_win.transient(self)
        about_win.grab_set()
        parent_x = self.winfo_x()
        parent_y = self.winfo_y()
        parent_width = self.winfo_width()
        parent_height = self.winfo_height()
        pos_x = parent_x + (parent_width // 2) - (dialog_width // 2)
        pos_y = parent_y + (parent_height // 2) - (dialog_height // 2)
        about_win.geometry(f"{dialog_width}x{dialog_height}+{pos_x}+{pos_y}")
        main_frame = ttk.Frame(about_win, padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        ttk.Label(main_frame, text="Zaawansowany Konwerter PDF", font=("Helvetica", 14, "bold")).pack(pady=(0,10))
        ttk.Label(main_frame, text="Wersja: 4.2.0 Enterprise Edition").pack(pady=3)
        ttk.Label(main_frame, text="Autor: Alan Steinbarth").pack(pady=3)
        ttk.Label(main_frame, text="Wykorzystuje biblioteki: pdf2docx, pdfminer.six, PyMuPDF, Pillow, (opcjonalnie) Pytesseract.").pack(pady=3)
        ttk.Label(main_frame, text="© 2025").pack(pady=3)
        # Instrukcja i linki
        help_textbox = scrolledtext.ScrolledText(main_frame, height=12, wrap=tk.WORD, state=tk.NORMAL)
        help_textbox.insert(tk.END, self.help_text)
        help_textbox.config(state=tk.DISABLED)
        help_textbox.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        close_button = ttk.Button(main_frame, text="Zamknij", command=about_win.destroy)
        close_button.pack(pady=(20,0))
        about_win.protocol("WM_DELETE_WINDOW", about_win.destroy)

    def center_window(self):
        """Centruje okno na ekranie"""
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        # Upewnijmy się, że width i height nie są zerowe (co może się zdarzyć na początku)
        if width <= 1: # Czasami winfo_width() może zwrócić 1 na starcie
            width = int(self.geometry().split('x')[0].split('+')[0]) # Próba odczydu z geometrii
        if height <= 1:
            height = int(self.geometry().split('x')[1].split('+')[0]) # Próba odczydu z geometrii

        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        pos_x = (screen_width // 2) - (width // 2)
        pos_y = (screen_height // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{pos_x}+{pos_y}")

    def on_closing(self):
        """Obsługuje zamykanie aplikacji."""
        if self.conversion_thread and self.conversion_thread.is_alive():
            if messagebox.askyesno("Potwierdzenie", "Trwa konwersja. Czy na pewno chcesz przerwać i zamknąć aplikację?"):
                logging.warning("Przerwanie konwersji i zamknięcie aplikacji przez użytkownika.")
                self.cancel_conversion(wait_for_thread=False) # Anuluj bez czekania, jeśli użytkownik chce szybko zamknąć
                self.destroy()
            else:
                return # Nie zamykaj
        else:
            self.destroy()

    def select_files(self):
        """Otwiera okno dialogowe do wyboru plików PDF i aktualizuje listę."""
        initial_dir = self.output_dir if os.path.isdir(self.output_dir) else os.path.expanduser("~")

        new_files = filedialog.askopenfilenames(
            title="Wybierz pliki PDF",
            initialdir=initial_dir,
            filetypes=(("Pliki PDF", "*.pdf"), ("Wszystkie pliki", "*.*"))
        )
        if new_files:
            for f in new_files:
                if f not in self.selected_files:
                    self.selected_files.append(f)
            self.update_files_listbox()
            if not self.files_listbox.curselection() and self.selected_files:
                self.files_listbox.selection_set(0) # Zaznacz pierwszy, jeśli lista niepusta
            self.on_file_select() # Zaktualizuj informacje o pliku

    def update_files_listbox(self):
        """Aktualizuje listę plików w Listbox."""
        self.files_listbox.delete(0, tk.END)
        for f_path in self.selected_files:
            self.files_listbox.insert(tk.END, os.path.basename(f_path))
        # Można dodać logikę aktualizacji statusu, np. liczby plików

    def add_folder(self):
        """Otwiera okno dialogowe do wyboru folderu i dodaje pliki PDF."""
        initial_dir = self.output_dir if os.path.isdir(self.output_dir) else os.path.expanduser("~")
        folder_path = filedialog.askdirectory(title="Wybierz folder z plikami PDF", initialdir=initial_dir)

        if folder_path:
            found_files = []
            for root, _, files in os.walk(folder_path):
                for file in files:
                    if file.lower().endswith(".pdf"):
                        full_path = os.path.join(root, file)
                        if full_path not in self.selected_files and full_path not in found_files:
                            found_files.append(full_path)

            if found_files:
                self.selected_files.extend(found_files)
                self.update_files_listbox()
                logging.info("Dodano %d plików PDF z folderu: %s", len(found_files), folder_path)
                if not self.files_listbox.curselection() and self.selected_files:
                    self.files_listbox.selection_set(0)
                self.on_file_select()
            else:
                logging.warning("Nie znaleziono plików PDF w folderze: %s", folder_path)

    def remove_selected_files(self):
        """Usuwa zaznaczone pliki z listy."""
        selected_indices = self.files_listbox.curselection()
        if not selected_indices:
            logging.warning("Nie wybrano plików do usunięcia.")
            return

        # Usuwamy od końca, aby uniknąć problemów z indeksami
        for index in reversed(selected_indices):
            del self.selected_files[index]

        self.update_files_listbox()
        logging.info("Usunięto %d plików z listy.", len(selected_indices))
        if self.selected_files and not self.files_listbox.curselection():
            self.files_listbox.selection_set(0) # Zaznacz pierwszy, jeśli coś zostało
        self.on_file_select() # Zaktualizuj info (wyczyści, jeśli lista pusta)

    def clear_file_list(self):
        """Czyści listę wybranych plików."""
        if not self.selected_files:
            logging.info("Lista plików jest już pusta.")
            return

        self.selected_files.clear()
        self.update_files_listbox()
        logging.info("Wyczyszczono listę plików.")
        self.on_file_select() # Wywoła _clear_file_info_display

    def select_output_dir(self):
        """Otwiera okno dialogowe do wyboru folderu docelowego."""
        initial_dir = self.output_dir if os.path.isdir(self.output_dir) else os.path.expanduser("~")
        new_dir = filedialog.askdirectory(title="Wybierz folder docelowy", initialdir=initial_dir)
        if new_dir:
            self.output_dir = new_dir
            self.output_dir_entry.config(state=tk.NORMAL)
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, self.output_dir)
            self.output_dir_entry.config(state="readonly")
            logging.info("Zmieniono folder docelowy na: %s", self.output_dir)

    def _toggle_controls_state(self, state):
        """Włącza lub wyłącza główne kontrolki."""
        widgets_to_toggle = [
            self.select_files_button, self.add_folder_button,
            self.remove_button, self.clear_list_button,
            self.select_output_dir_button, self.docx_radio, self.txt_radio
        ]
        for widget in widgets_to_toggle:
            if widget:
                widget.config(state=state)

    def start_conversion(self):
        """Rozpoczyna proces konwersji plików w osobnym wątku."""
        # print("DEBUG: start_conversion called") # DEBUG
        if not self.selected_files:
            logging.warning("Nie wybrano plików do konwersji.")
            messagebox.showwarning("Brak plików", "Proszę wybrać pliki do konwersji.")
            return

        if self.conversion_thread and self.conversion_thread.is_alive():
            logging.warning("Konwersja jest już w toku.")
            messagebox.showwarning("Konwersja w toku", "Poczekaj na zakończenie bieżącej konwersji.")
            return

        # print(f"DEBUG: Starting conversion for {len(self.selected_files)} files. Output format: {self.output_format.get()}") # DEBUG
        logging.info("Rozpoczynanie konwersji %d plików do formatu %s...", len(self.selected_files), self.output_format.get().upper())
        self._toggle_controls_state(tk.DISABLED)
        self.convert_button.config(state=tk.DISABLED)
        self.cancel_button.config(state=tk.NORMAL)

        self.cancel_event.clear()
        self.progress_bar["value"] = 0
        self.progress_bar["maximum"] = len(self.selected_files)
        # print(f"DEBUG: Progress bar max set to {self.progress_bar['maximum']}") # DEBUG

        self.conversion_thread = threading.Thread(target=self._perform_conversion_threaded, daemon=True)
        self.conversion_thread.start()
        # print("DEBUG: Conversion thread started") # DEBUG

    def _convert_single_file(self, file_path):
        base_name = os.path.basename(file_path)
        name_without_ext = os.path.splitext(base_name)[0]
        output_format_val = self.output_format.get()
        output_file_name = f"{name_without_ext}.{output_format_val}"
        output_file_path = os.path.join(self.output_dir, output_file_name)
        try:
            text_content = extract_text(file_path)
            # Jeśli tekst jest pusty lub bardzo krótki, automatycznie uruchom OCR
            if not text_content or len(text_content.strip()) < 20:
                logging.info("Automatyczne uruchamianie OCR dla '%s' (brak tekstu w PDF)...", base_name)
                ocr_texts = []
                ocr_pages_with_text = 0
                if TESSERACT_AVAILABLE and pytesseract:
                    try:
                        from PIL import ImageOps, ImageFilter
                        doc = fitz.open(file_path)
                        for page_num in range(len(doc)):
                            page = doc[page_num]
                            try:
                                pix = self._get_pixmap_compat(page, fitz.Matrix(3, 3))
                            except (OSError, IOError, ValueError, RuntimeError) as e:
                                logging.error("Błąd generowania obrazu do OCR: %s", e)
                                continue
                            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                            img = img.convert("L")
                            img = ImageOps.autocontrast(img)
                            img = img.filter(ImageFilter.SHARPEN)
                            threshold = 150
                            table = [0 if i < threshold else 255 for i in range(256)]
                            img = img.point(table, 'L')
                            ocr_text = pytesseract.image_to_string(img, lang=self.ocr_lang, config='--psm 6')
                            ocr_text = ocr_text.strip()
                            if ocr_text:
                                ocr_pages_with_text += 1
                            ocr_texts.append(ocr_text)
                            pix = None
                        doc.close()
                        text_content = "\n\n--- Koniec strony ---\n\n".join(ocr_texts)
                        logging.info("OCR zakończony: %d/%d stron zawierało tekst.", ocr_pages_with_text, len(ocr_texts))
                    except (OSError, IOError, ValueError, RuntimeError) as ocr_err:
                        logging.error("Błąd OCR: %s. Nie udało się wyciągnąć tekstu z obrazu.", ocr_err)
                        text_content = ""
                else:
                    logging.warning("OCR niedostępny – nie można przetworzyć pliku graficznego.")
                    text_content = ""
            if output_format_val == "docx":
                from docx import Document
                docx_doc = Document()
                if text_content.strip():
                    for page_text in text_content.split("\n\n--- Koniec strony ---\n\n"):
                        lines = [line for line in page_text.splitlines() if line.strip()]
                        if lines:
                            for line in lines:
                                docx_doc.add_paragraph(line)
                            docx_doc.add_paragraph("")
                    # Usuń ostatni pusty akapit jeśli jest
                    if docx_doc.paragraphs and not docx_doc.paragraphs[-1].text.strip():
                        p = docx_doc.paragraphs[-1]._element  # pylint: disable=protected-access
                        p.getparent().remove(p)
                    docx_doc.save(output_file_path)
                    logging.info("Pomyślnie zapisano '%s' jako '%s'.", base_name, output_file_name)
                    return True
                
                logging.warning("Nie udało się wyodrębnić tekstu z pliku '%s'. Plik DOCX nie został utworzony.", base_name)
                return False
            if output_format_val == "txt":
                if text_content.strip():
                    with open(output_file_path, "w", encoding="utf-8") as f:
                        f.write(text_content)
                    logging.info("Pomyślnie zapisano '%s' jako '%s'.", base_name, output_file_name)
                    return True
                
                logging.warning("Nie udało się wyodrębnić tekstu z pliku '%s'. Plik TXT nie został utworzony.", base_name)
                return False
        except (OSError, IOError, ValueError, RuntimeError) as e:
            logging.error("Błąd podczas konwersji pliku '%s': %s", base_name, e)
            return False

    # =========================================================================
    # 🔄 LOGIKA KONWERSJI I PRZETWARZANIA
    # =========================================================================

    def _perform_conversion_threaded(self):
        """
        ## 🔄 Główna Logika Konwersji (Threading)
        
        **Opis**: Wykonuje konwersję plików PDF w osobnym wątku, 
        aby nie blokować interfejsu użytkownika.
        
        **Proces konwersji**:
        1. 📂 Iteracja przez wybrane pliki
        2. 📄 Przetwarzanie każdego pliku PDF
        3. 🔍 Ekstrakcja tekstu (pdfminer)
        4. 🖼️ OCR dla skanów (jeśli potrzebne)
        5. 💾 Zapis do DOCX/TXT
        6. 📊 Aktualizacja progress bar
        
        **Threading safety**:
        - Używa `self.cancel_event` do przerwania
        - Thread-safe komunikacja z GUI przez kolejki
        - Graceful cleanup w przypadku błędów
        
        **Error handling**: Kontynuuje przetwarzanie mimo błędów w pojedynczych plikach
        """
        # print("DEBUG: _perform_conversion_threaded started") # DEBUG
        files_converted_count = 0
        try:
            for i, file_path in enumerate(self.selected_files):
                if self.cancel_event.is_set():
                    logging.warning("Konwersja przerwana przez użytkownika.")
                    break
                logging.info("Przetwarzanie pliku (%d/%d): %s", i+1, len(self.selected_files), os.path.basename(file_path))
                try:
                    success = self._convert_single_file(file_path)
                except (OSError, IOError, ValueError, RuntimeError) as file_exc:
                    logging.error("Błąd krytyczny podczas konwersji pliku %s: %s", file_path, file_exc)
                    success = False
                if success:
                    files_converted_count += 1
                else:
                    logging.warning("Pominięto plik %s z powodu błędu.", os.path.basename(file_path))
                current_progress = i + 1
                self.after(0, lambda val=current_progress: self.progress_bar.config(value=val))
        except (OSError, IOError, ValueError, RuntimeError, threading.ThreadError) as e:
            logging.error("Wystąpił krytyczny błąd w głównym wątku konwersji: %s", e)
        finally:
            self.after(0, self._finalize_conversion)

    def cancel_conversion(self, wait_for_thread=True):
        """Anuluje trwającą konwersję."""
        if self.conversion_thread and self.conversion_thread.is_alive():
            if not self.cancel_event.is_set():
                self.cancel_event.set()
                logging.info("Anulowanie konwersji...")
                if wait_for_thread:
                    self.conversion_thread.join(timeout=2)
                self._toggle_controls_state(tk.NORMAL)
                self.convert_button.config(state=tk.NORMAL)
                self.cancel_button.config(state=tk.DISABLED)
            else:
                logging.info("Anulowanie jest już w toku.")
        else:
            logging.info("Brak aktywnej konwersji do anulowania.")
            self.cancel_button.config(state=tk.DISABLED)

    def _finalize_conversion(self):
        self._toggle_controls_state(tk.NORMAL)
        self.convert_button.config(state=tk.NORMAL)
        self.cancel_button.config(state=tk.DISABLED)
        logging.info("Konwersja zakończona.")

    def _get_pdf_page_count(self, pdf_path):
        """Pobiera liczbę stron z pliku PDF."""
        try:
            count = 0
            for _ in extract_pages(pdf_path):
                count += 1
            return count
        except (OSError, IOError, ValueError, RuntimeError) as e:
            logging.warning("Nie udało się pobrać liczby stron dla '%s': %s. Szczegóły w logach konsoli, jeśli dostępne.", os.path.basename(pdf_path), type(e).__name__)
            return None

    def on_file_select(self, _event=None):
        """Obsługuje zdarzenie wyboru pliku na liście i aktualizuje informacje oraz podgląd."""
        if not hasattr(self, 'files_listbox') or not hasattr(self, 'selected_files'):
            return

        selected_indices = self.files_listbox.curselection()
        if not selected_indices:
            self._clear_file_info_display()
            return
        selected_index = selected_indices[0]
        if 0 <= selected_index < len(self.selected_files):
            file_path = self.selected_files[selected_index]
            self._display_file_info(file_path)
            self._preview_pdf_page(file_path)
            self.preview_canvas.update_idletasks()
        else:
            self._clear_file_info_display()
            self._clear_pdf_preview()
            logging.debug("Niespójność indeksu w on_file_select: %d vs len %d", selected_index, len(self.selected_files))

    def _display_file_info(self, file_path):
        """Wyświetla informacje o podanym pliku PDF."""
        if not hasattr(self, 'info_file_name_var'):
            return

        if not os.path.exists(file_path):
            self._clear_file_info_display()
            logging.warning("Plik nie istnieje w _display_file_info: %s", file_path)
            return

        file_name = os.path.basename(file_path)
        size_str = "N/A" # Inicjalizacja

        try:
            file_size_bytes = os.path.getsize(file_path)
            if file_size_bytes < 1024:
                size_str = f"{file_size_bytes} B"
            elif file_size_bytes < 1024 * 1024:
                size_str = f"{file_size_bytes / 1024:.2f} KB"
            else:
                size_str = f"{file_size_bytes / (1024 * 1024):.2f} MB"
        except OSError as e:
            logging.warning("Nie udało się pobrać rozmiaru pliku '%s': %s", file_name, e)
            size_str = "N/A"

        page_count = self._get_pdf_page_count(file_path)
        page_count_str = str(page_count) if page_count is not None else "N/A"

        self.info_file_name_var.set(f"{file_name}")
        self.info_file_size_var.set(f"{size_str}")
        self.info_file_pages_var.set(f"{page_count_str}")

    def _clear_file_info_display(self):
        """Czyści wyświetlane informacje o pliku."""
        if hasattr(self, 'info_file_name_var'):
            self.info_file_name_var.set("-")
            self.info_file_size_var.set("-")
            self.info_file_pages_var.set("-")
            self._clear_pdf_preview() # Wyczyść również podgląd

    def _clear_pdf_preview(self):
        """Czyści canvas podglądu PDF."""
        if hasattr(self, 'preview_canvas'):
            self.preview_canvas.delete("all")
            self.current_preview_image = None

    def _get_pixmap_compat(self, page, matrix):
        """Zwraca pixmapę strony PDF, obsługując różne wersje PyMuPDF."""
        if hasattr(page, 'get_pixmap'):
            return page.get_pixmap(matrix=matrix)
        if hasattr(page, 'getPixmap'):
            return page.getPixmap(matrix=matrix)
        
        raise RuntimeError("Brak kompatybilnej metody get_pixmap/getPixmap w obiekcie Page.")

    def _preview_pdf_page(self, pdf_path, page_num=0):
        if not hasattr(self, 'preview_canvas'):
            return
        self._clear_pdf_preview()
        try:
            doc = fitz.open(pdf_path)
            if not doc:
                logging.error("Nie można otworzyć pliku PDF: %s", pdf_path)
                return
            if page_num >= len(doc):
                logging.error("Nieprawidłowy numer strony: %d", page_num)
                doc.close()
                return
            page = doc[page_num]
            preview_width = self.preview_canvas.winfo_width()
            preview_height = self.preview_canvas.winfo_height()
            if preview_width <= 1 or preview_height <= 1:
                # macOS-safe: ponów próbę po 200 ms
                self.after(200, lambda: self._preview_pdf_page(pdf_path, page_num))
                doc.close()
                return
            scale = min(preview_width / page.rect.width, preview_height / page.rect.height)
            mat = fitz.Matrix(scale, scale)
            try:
                pix = self._get_pixmap_compat(page, mat)
            except (OSError, IOError, ValueError, RuntimeError) as e:
                logging.error("Błąd generowania podglądu PDF: %s", e)
                doc.close()
                return
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            self.current_preview_image = ImageTk.PhotoImage(img)
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(preview_width/2, preview_height/2, image=self.current_preview_image, anchor=tk.CENTER)
            pix = None
            doc.close()
        except (OSError, IOError, ValueError, RuntimeError) as e:
            logging.error("Błąd podczas generowania podglądu PDF: %s", e)
            self._clear_pdf_preview()

    def add_theme_switcher(self):
        """Add theme switcher button to the main window."""
        # Przycisk przełączania motywu w górnej części okna - nowoczesny design
        c = self.theme_colors[self.theme_mode]
        self.theme_switch_btn = tk.Button(
            self,
            text="🌙" if self.theme_mode == "light" else "☀️",
            command=self.toggle_theme,
            bg=c["button_secondary"],
            fg=c["button_secondary_text"],
            relief=tk.FLAT,
            font=("Helvetica", 14, "bold"),
            bd=1,
            highlightthickness=1,
            highlightbackground=c["border"],
            activebackground=c["button_secondary_active"],
            activeforeground=c["button_secondary_text"],
            width=3,
            height=1
        )
        self.theme_switch_btn.place(x=10, y=10, width=40, height=32)

    def toggle_theme(self):
        """Toggle between light and dark theme."""
        self.theme_mode = "dark" if self.theme_mode == "light" else "light"
        self.update_theme()

    def update_theme(self):
        """Update all UI elements with current theme colors."""
        c = self.theme_colors[self.theme_mode]

        # Tło główne
        self.configure(bg=c["bg"])
        self.main_frame.configure(bg=c["bg"])

        # Panele
        self.left_panel_frame.configure(bg=c["panel"], highlightbackground=c["border"], highlightthickness=1)
        self.right_panel_frame.configure(bg=c["panel"], highlightbackground=c["border"], highlightthickness=1)

        # Lista plików z lepszym kontrastem
        self.files_listbox.configure(
            bg=c["listbox"],
            fg=c["label"],
            selectbackground=c["accent"],
            selectforeground="white",
            highlightbackground=c["border"],
            highlightcolor=c["accent"]
        )
        self.files_listbox_frame.configure(bg=c["panel"])
        self.file_buttons_frame.configure(bg=c["panel"])

        # Ramka informacji o pliku
        self.file_info_frame.configure(bg=c["listbox"], fg=c["label"])

        # Ustawienia wyjściowe
        self.output_settings_frame.configure(bg=c["panel"])
        self.output_dir_label.configure(bg=c["panel"], fg=c["label"], font=("Helvetica", 11, "bold"))
        self.output_dir_entry_frame.configure(bg=c["panel"])
        self.output_dir_entry.configure(
            bg=c["entry"],
            fg=c["label"],
            highlightbackground=c["entry_border"],
            highlightcolor=c["accent"],
            insertbackground=c["label"],
            readonlybackground=c["entry"],
            disabledbackground=c["entry"],
            disabledforeground=c["label"]
        )

        # Przycisk zmiany folderu - styl secondary
        self.select_output_dir_button.configure(
            bg=c["button_secondary"],
            fg=c["button_secondary_text"],
            activebackground=c["button_secondary_active"],
            font=("Helvetica", 10, "bold"),
            relief=tk.FLAT,
            bd=1,
            highlightbackground=c["border"]
        )

        # Ramka formatów
        self.format_frame.configure(bg=c["listbox"], fg=c["label"])
        self.docx_radio.configure(
            bg=c["listbox"],
            fg=c["label"],
            selectcolor=c["accent"],
            activebackground=c["listbox"],
            font=("Helvetica", 11)
        )
        self.txt_radio.configure(
            bg=c["listbox"],
            fg=c["label"],
            selectcolor=c["accent"],
            activebackground=c["listbox"],
            font=("Helvetica", 11)
        )

        # Pasek postępu
        self.progress_bar.configure(style="Custom.Horizontal.TProgressbar")

        # Ramka przycisków akcji
        self.action_buttons_frame.configure(bg=c["panel"])

        # Przycisk konwertuj - jednolity styl jak wszystkie inne przyciski
        self.convert_button.configure(
            bg=c["button_secondary"],
            fg=c["button_secondary_text"],
            activebackground=c["button_secondary_active"],
            font=("Helvetica", 11, "bold"),
            relief=tk.FLAT,
            bd=1,
            highlightthickness=1,
            highlightbackground=c["border"]
        )

        # Przycisk anuluj - jednolity styl ale z wyraźniejszym kontrastem
        self.cancel_button.configure(
            bg=c["button_secondary"],
            fg=c["label"],  # Zmieniono na label zamiast button_secondary_text dla lepszego kontrastu
            activebackground=c["button_secondary_active"],
            activeforeground=c["label"],
            font=("Helvetica", 11, "bold"),
            relief=tk.FLAT,
            bd=1,
            highlightthickness=1,
            highlightbackground=c["border"]
        )

        # Przycisk pomocy - styl secondary
        self.help_button.configure(
            bg=c["button_secondary"],
            fg=c["button_secondary_text"],
            activebackground=c["button_secondary_active"],
            font=("Helvetica", 11, "bold"),
            relief=tk.FLAT,
            bd=1,
            highlightbackground=c["border"]
        )

        # Stopka
        self.footer_label.configure(bg=c["footer"], fg=c["footer_fg"])

        # Canvas podglądu
        self.preview_canvas.configure(bg=c["listbox"], highlightbackground=c["border"], highlightthickness=1)

        # Widget logów - aktualizacja kolorów
        if hasattr(self, 'log_text_widget'):
            self.log_text_widget.configure(
                bg=c["listbox"],
                fg=c["label"],
                insertbackground=c["label"],
                selectbackground=c["accent"],
                selectforeground="white"
            )

        # Aktualizuj wszystkie przyciski zarządzania plikami
        for btn in [self.select_files_button, self.add_folder_button, self.remove_button, self.clear_list_button]:
            btn.configure(
                bg=c["button_secondary"],
                fg=c["button_secondary_text"],
                activebackground=c["button_secondary_active"],
                font=("Helvetica", 10, "bold"),
                relief=tk.FLAT,
                bd=1,
                highlightbackground=c["border"]
            )

        # Etykiety w ramkach
        for widget in self.file_info_frame.winfo_children():
            if isinstance(widget, tk.Label):
                widget.configure(bg=c["listbox"], fg=c["label"], font=("Helvetica", 10))

        # Przełącznik motywu
        if self.theme_switch_btn is not None:
            self.theme_switch_btn.configure(
                bg=c["button_secondary"],
                fg=c["button_secondary_text"],
                activebackground=c["button_secondary_active"],
                highlightbackground=c["border"]
            )
            self.theme_switch_btn.config(text="☀️" if self.theme_mode == "dark" else "🌙")

        # Style ttk
        style = ttk.Style()
        style.theme_use('clam')
        style.configure(
            "Custom.Horizontal.TProgressbar",
            background=c["progress"],
            troughcolor=c["progress_trough"],
            bordercolor=c["border"],
            lightcolor=c["progress"],
            darkcolor=c["progress"],
            borderwidth=1
        )

        # Odświeżenie całego okna
        self.update_idletasks()

    def create_widgets(self):
        """Create and arrange all GUI widgets."""
        # Główny kontener (wszystko poza stopką)
        self.configure(bg=self.theme_colors[self.theme_mode]["bg"])
        self.main_frame = tk.Frame(self, bg=self.theme_colors[self.theme_mode]["bg"])
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=16, pady=12)
        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)
        self.main_frame.rowconfigure(0, weight=1)
        self.main_frame.columnconfigure(0, weight=1)

        # Dodaj przycisk przełączania motywu dzień/noc
        self.add_theme_switcher()

        # Lewy panel (Kontrolki i Lista plików)
        self.left_panel_frame = tk.Frame(self.main_frame, bg=self.theme_colors[self.theme_mode]["panel"], bd=0, highlightthickness=0)
        self.left_panel_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0,18), pady=0)
        self.left_panel_frame.pack_propagate(False)
        self.left_panel_frame.rowconfigure(1, weight=2)
        self.left_panel_frame.columnconfigure(0, weight=1)

        # Przycisk Pomoc/FAQ NA GÓRZE
        self.help_button = tk.Button(self.left_panel_frame, text="Pomoc / O programie", command=self.show_about_dialog, bg=self.theme_colors[self.theme_mode]["button"], fg=self.theme_colors[self.theme_mode]["label"], relief=tk.FLAT, font=("Helvetica", 11, "bold"), bd=0, highlightthickness=0, activebackground=self.theme_colors[self.theme_mode]["button_active"])
        self.help_button.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        # Lista wybranych plików
        self.files_listbox_frame = tk.Frame(self.left_panel_frame, bg=self.theme_colors[self.theme_mode]["panel"])
        self.files_listbox_frame.grid(row=1, column=0, sticky="nsew")
        self.files_listbox = tk.Listbox(self.files_listbox_frame, selectmode=tk.EXTENDED, exportselection=False, bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], relief=tk.FLAT, font=("Helvetica", 10), highlightthickness=0, bd=0)
        self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.files_listbox.bind('<<ListboxSelect>>', self.on_file_select)
        self.files_scrollbar = tk.Scrollbar(self.files_listbox_frame, orient=tk.VERTICAL, command=self.files_listbox.yview)
        self.files_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.files_listbox.config(yscrollcommand=self.files_scrollbar.set)

        # Przyciski zarządzania plikami pod listą - nowoczesny styl
        self.file_buttons_frame = tk.Frame(self.left_panel_frame, bg=self.theme_colors[self.theme_mode]["panel"])
        self.file_buttons_frame.grid(row=2, column=0, sticky="ew", pady=(2, 0))
        btn_style = {
            "bg": self.theme_colors[self.theme_mode]["button_secondary"],
            "fg": self.theme_colors[self.theme_mode]["button_secondary_text"],
            "activebackground": self.theme_colors[self.theme_mode]["button_secondary_active"],
            "relief": tk.FLAT,
            "font": ("Helvetica", 10, "bold"),
            "bd": 1,
            "highlightthickness": 1,
            "highlightbackground": self.theme_colors[self.theme_mode]["border"]
        }
        self.select_files_button = tk.Button(self.file_buttons_frame, text="Wybierz pliki", command=self.select_files, width=12, **btn_style)
        self.select_files_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0,2))
        self.add_folder_button = tk.Button(self.file_buttons_frame, text="Dodaj folder", command=self.add_folder, width=12, **btn_style)
        self.add_folder_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2,0))
        self.remove_button = tk.Button(self.file_buttons_frame, text="Usuń wybrane", command=self.remove_selected_files, width=12, **btn_style)
        self.remove_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0,2))
        self.clear_list_button = tk.Button(self.file_buttons_frame, text="Wyczyść listę", command=self.clear_file_list, width=12, **btn_style)
        self.clear_list_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2,0))

        # Sekcja informacji o pliku
        self.file_info_frame = tk.LabelFrame(self.left_panel_frame, text="Informacje o pliku", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], padx=10, pady=5, font=("Helvetica", 10, "bold"), bd=1, relief=tk.GROOVE, highlightbackground=self.theme_colors[self.theme_mode]["button"], highlightcolor=self.theme_colors[self.theme_mode]["button"])
        self.file_info_frame.grid(row=3, column=0, sticky="ew", pady=(10, 5))
        tk.Label(self.file_info_frame, text="Nazwa:", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"]).grid(row=0, column=0, sticky=tk.W, padx=(0,5), pady=1)
        self.info_file_name_var = tk.StringVar(value="-")
        self.info_file_name_label = tk.Label(self.file_info_frame, textvariable=self.info_file_name_var, anchor=tk.W, wraplength=200, bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10))
        self.info_file_name_label.grid(row=0, column=1, sticky=tk.EW, pady=1)
        tk.Label(self.file_info_frame, text="Rozmiar:", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"]).grid(row=1, column=0, sticky=tk.W, padx=(0,5), pady=1)
        self.info_file_size_var = tk.StringVar(value="-")
        self.info_file_size_label = tk.Label(self.file_info_frame, textvariable=self.info_file_size_var, anchor=tk.W, bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10))
        self.info_file_size_label.grid(row=1, column=1, sticky=tk.EW, pady=1)
        tk.Label(self.file_info_frame, text="Liczba stron:", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"]).grid(row=2, column=0, sticky=tk.W, padx=(0,5), pady=1)
        self.info_file_pages_var = tk.StringVar(value="-")
        self.info_file_pages_label = tk.Label(self.file_info_frame, textvariable=self.info_file_pages_var, anchor=tk.W, bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10))
        self.info_file_pages_label.grid(row=2, column=1, sticky=tk.EW, pady=1)
        self.file_info_frame.columnconfigure(1, weight=1)

        # Sekcja folderu docelowego i formatu
        self.output_settings_frame = tk.Frame(self.left_panel_frame, bg=self.theme_colors[self.theme_mode]["panel"])
        self.output_settings_frame.grid(row=4, column=0, sticky="ew", pady=(10,5))
        self.output_dir_label = tk.Label(self.output_settings_frame, text="Folder docelowy:", bg=self.theme_colors[self.theme_mode]["panel"], fg=self.theme_colors[self.theme_mode]["label"])
        self.output_dir_label.pack(side=tk.TOP, anchor=tk.W)
        self.output_dir_entry_frame = tk.Frame(self.output_settings_frame, bg=self.theme_colors[self.theme_mode]["panel"])
        self.output_dir_entry_frame.pack(fill=tk.X)
        self.output_dir_entry = tk.Entry(
            self.output_dir_entry_frame,
            state="readonly",
            bg=self.theme_colors[self.theme_mode]["entry"],
            fg=self.theme_colors[self.theme_mode]["label"],
            relief=tk.FLAT,
            font=("Helvetica", 10),
            readonlybackground=self.theme_colors[self.theme_mode]["entry"],
            disabledbackground=self.theme_colors[self.theme_mode]["entry"],
            disabledforeground=self.theme_colors[self.theme_mode]["label"]
        )
        self.output_dir_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5))
        self.output_dir_entry.config(state=tk.NORMAL)
        self.output_dir_entry.delete(0, tk.END)
        self.output_dir_entry.insert(0, self.output_dir)
        self.output_dir_entry.config(state="readonly")
        self.select_output_dir_button = tk.Button(self.output_dir_entry_frame, text="Zmień", command=self.select_output_dir, width=8, bg=self.theme_colors[self.theme_mode]["button"], fg=self.theme_colors[self.theme_mode]["label"], relief=tk.FLAT, font=("Helvetica", 10), bd=0, highlightthickness=0, activebackground=self.theme_colors[self.theme_mode]["button_active"])
        self.select_output_dir_button.pack(side=tk.LEFT)
        self.format_frame = tk.LabelFrame(self.output_settings_frame, text="Format wyjściowy", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], padx=5, font=("Helvetica", 10, "bold"), bd=1, relief=tk.GROOVE, highlightbackground=self.theme_colors[self.theme_mode]["button"], highlightcolor=self.theme_colors[self.theme_mode]["button"])
        self.format_frame.pack(fill=tk.X, pady=(10,5), anchor=tk.W)
        self.docx_radio = tk.Radiobutton(self.format_frame, text="DOCX", variable=self.output_format, value="docx", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], selectcolor=self.theme_colors[self.theme_mode]["button"], font=("Helvetica", 10))
        self.docx_radio.pack(side=tk.LEFT, padx=(0,10), expand=True)
        self.txt_radio = tk.Radiobutton(self.format_frame, text="TXT", variable=self.output_format, value="txt", bg=self.theme_colors[self.theme_mode]["listbox"], fg=self.theme_colors[self.theme_mode]["label"], selectcolor=self.theme_colors[self.theme_mode]["button"], font=("Helvetica", 10))
        self.txt_radio.pack(side=tk.LEFT, padx=(0,10), expand=True)

        # Pasek postępu
        self.progress_bar = ttk.Progressbar(self.left_panel_frame, orient="horizontal", mode="determinate")
        self.progress_bar.grid(row=5, column=0, sticky="ew", pady=(10, 5))

        # Przyciski akcji na dole lewego panelu - nowoczesny styl
        self.action_buttons_frame = tk.Frame(self.left_panel_frame, bg=self.theme_colors[self.theme_mode]["panel"])
        self.action_buttons_frame.grid(row=6, column=0, sticky="ew", pady=(5,0))

        # Przycisk konwertuj - główny przycisk (primary) - jednolity styl jak inne przyciski
        self.convert_button = tk.Button(
            self.action_buttons_frame,
            text="Konwertuj",
            command=self.start_conversion,
            width=12,
            bg=self.theme_colors[self.theme_mode]["button_secondary"],
            fg=self.theme_colors[self.theme_mode]["button_secondary_text"],
            activebackground=self.theme_colors[self.theme_mode]["button_secondary_active"],
            relief=tk.FLAT,
            font=("Helvetica", 11, "bold"),
            bd=1,
            highlightthickness=1,
            highlightbackground=self.theme_colors[self.theme_mode]["border"]
        )
        self.convert_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0,2))

        # Przycisk anuluj - drugi przycisk (secondary) - lepszy kontrast
        self.cancel_button = tk.Button(
            self.action_buttons_frame,
            text="Anuluj",
            command=self.cancel_conversion,
            state=tk.DISABLED,
            width=12,
            bg=self.theme_colors[self.theme_mode]["button_secondary"],
            fg=self.theme_colors[self.theme_mode]["label"],  # Zmieniono na label dla lepszego kontrastu
            activebackground=self.theme_colors[self.theme_mode]["button_secondary_active"],
            activeforeground=self.theme_colors[self.theme_mode]["label"],
            relief=tk.FLAT,
            font=("Helvetica", 11, "bold"),
            bd=1,
            highlightthickness=1,
            highlightbackground=self.theme_colors[self.theme_mode]["border"]
        )
        self.cancel_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2,0))

        # Prawy panel (podgląd PDF)
        self.right_panel_frame = tk.Frame(self.main_frame, bg=self.theme_colors[self.theme_mode]["panel"], bd=0, highlightthickness=0)
        self.right_panel_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.right_panel_frame.rowconfigure(0, weight=1)
        self.right_panel_frame.columnconfigure(0, weight=1)
        self.preview_canvas = tk.Canvas(self.right_panel_frame, bg=self.theme_colors[self.theme_mode]["listbox"], height=320, highlightthickness=0, bd=0)
        self.preview_canvas.grid(row=0, column=0, sticky="nsew", padx=10, pady=(0, 5))

        # Stopka z informacją o autorze
        self.footer_label = tk.Label(self, text="Autor: Alan Steinbarth | © 2025", anchor=tk.CENTER, font=("Helvetica", 10, "italic"), bg=self.theme_colors[self.theme_mode]["footer"], fg=self.theme_colors[self.theme_mode]["footer_fg"])
        self.footer_label.pack(side=tk.BOTTOM, fill=tk.X, pady=(2, 2))

        # --- Style dla jasnego tła ---
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TFrame", background=self.theme_colors[self.theme_mode]["bg"])
        style.configure("TLabel", background=self.theme_colors[self.theme_mode]["bg"], foreground=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10))
        style.configure("TButton", background=self.theme_colors[self.theme_mode]["button"], foreground=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10))
        style.configure("TLabelframe", background=self.theme_colors[self.theme_mode]["listbox"], foreground=self.theme_colors[self.theme_mode]["label"], font=("Helvetica", 10, "bold"))
        style.configure("TLabelframe.Label", background=self.theme_colors[self.theme_mode]["listbox"], foreground=self.theme_colors[self.theme_mode]["label"])
        style.configure("TPanedwindow", background=self.theme_colors[self.theme_mode]["bg"])
        style.configure("TProgressbar", background=self.theme_colors[self.theme_mode]["progress"], troughcolor=self.theme_colors[self.theme_mode]["progress_trough"], bordercolor=self.theme_colors[self.theme_mode]["progress_trough"], lightcolor=self.theme_colors[self.theme_mode]["progress"], darkcolor=self.theme_colors[self.theme_mode]["progress"])

        # Zastosuj motyw do wszystkich elementów
        self.update_theme()

if __name__ == "__main__":
    print("=== STARTING PDFtoDocxConverterApp ===")
    try:
        app = PDFtoDocxConverterApp()
        app.update_idletasks()
        app.deiconify()
        app.lift()
        app.focus_force()
        app.geometry("1024x768+100+100")
        app.mainloop()
    except Exception as e:  # pylint: disable=broad-exception-caught
        print(f"[FATAL ERROR] Aplikacja nie uruchomiła się! Szczegóły: {e}")
        traceback.print_exc()
    print("=== KONIEC PROGRAMU ===")

# =============================================================================
# 🚀 PUNKT WEJŚCIA APLIKACJI
# =============================================================================

if __name__ == "__main__":
    # =========================================================================
    # 🚀 Main Entry Point
    # 
    # Opis: Główny punkt wejścia aplikacji PDF Converter.
    # 
    # Proces uruchomienia:
    # 1. 📋 Wyświetlenie informacji o aplikacji
    # 2. ⚙️ Ładowanie konfiguracji z YAML
    # 3. 🏗️ Inicjalizacja głównego okna aplikacji
    # 4. 🖥️ Uruchomienie głównej pętli GUI
    # 5. 🛡️ Obsługa błędów krytycznych
    # 
    # Wymagania systemowe:
    # - Python 3.9+
    # - Tesseract OCR (opcjonalnie)
    # - Poppler utils
    # - Wymagane biblioteki z requirements.txt
    # 
    # Konfiguracja:
    # - Automatyczne wykrywanie motywu systemowego
    # - Domyślne rozmiary okna: 1024x768
    # - Pozycja startowa: +100+100
    # =========================================================================
    
    print("=== PDF to DOCX/TXT Converter v4.2.0 Enterprise Edition ===")
    print("Aplikacja uruchamia się...")
    
    try:
        app = PDFtoDocxConverterApp()
        app.lift()
        app.focus_force()
        app.geometry("1024x768+100+100")
        app.mainloop()
    except Exception as e:  # pylint: disable=broad-exception-caught
        print(f"[FATAL ERROR] Aplikacja nie uruchomiła się! Szczegóły: {e}")
        traceback.print_exc()
    print("=== KONIEC PROGRAMU ===")